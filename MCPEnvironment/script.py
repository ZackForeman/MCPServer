# pip install python-docx

from docx import Document

# Load existing document or create new one
try:
    doc = Document('output.docx')
except Exception:
    doc = Document()

# Add the desired phrase
p = doc.add_paragraph('this was added')

# Save the changes
doc.save('output.docx')